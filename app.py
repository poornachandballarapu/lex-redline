import streamlit as st
from pypdf import PdfReader
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from redlines import Redlines
import re
from io import BytesIO
from bs4 import BeautifulSoup
import markdown # We need this to standardise the input to HTML

# --- 1. CONFIGURATION ---
st.set_page_config(
    page_title="LexRedline | Universal Engine",
    page_icon="⚖️",
    layout="wide"
)

# --- 2. CSS STYLING ---
st.markdown("""
<style>
    .diff-box {
        background-color: white;
        padding: 50px;
        border: 1px solid #e0e0e0;
        border-radius: 4px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        font-family: 'Georgia', serif;
        font-size: 17px;
        line-height: 1.8;
        color: #2c3e50;
    }
    ins, .diff-add { background-color: #e3fcef; color: #155724; text-decoration: none; border-bottom: 2px solid #28a745; padding: 0 2px; }
    del, .diff-del { background-color: #fbeaea; color: #8a1f1f; text-decoration: line-through; opacity: 0.8; padding: 0 2px; }
</style>
""", unsafe_allow_html=True)

# --- 3. HELPER FUNCTIONS ---

def get_clean_text(uploaded_file):
    text = ""
    try:
        if uploaded_file.name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif uploaded_file.name.endswith(".docx"):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n\n"
        else:
            text = uploaded_file.getvalue().decode("utf-8")
        return text
    except Exception as e:
        return ""

def compare_html(text1, text2):
    marker = " BREAKTOKEN " 
    t1_clean = re.sub(r'\n+', marker, text1)
    t2_clean = re.sub(r'\n+', marker, text2)
    
    diff = Redlines(t1_clean, t2_clean)
    
    # Redlines outputs Markdown (e.g. **bold**, ~~strike~~)
    # We convert that Markdown to real HTML so BeautifulSoup can read it.
    md_output = diff.output_markdown
    html_output = markdown.markdown(md_output)
    
    # For web display, we just replace the token
    display_html = html_output.replace("BREAKTOKEN", "<br><br>")
    
    return html_output, display_html

def create_docx(html_content):
    """
    Universal Parser: Walks through the HTML tree.
    Detects <ins>, <del>, AND <span style="..."> tags.
    """
    doc = Document()
    doc.add_heading('Legal Redline Report', 0)
    
    # We use a dummy wrapper to ensure BeautifulSoup parses fragments correctly
    soup = BeautifulSoup(html_content, "html.parser")
    
    current_para = doc.add_paragraph()

    # Recursive function to walk through tags
    def process_element(element):
        nonlocal current_para
        
        # 1. If it's Text, write it
        if element.name is None:
            text = str(element)
            
            # Handle Paragraph Splits inside the text
            parts = text.split("BREAKTOKEN")
            for i, part in enumerate(parts):
                clean_part = part
                
                if clean_part:
                    # Determine style based on PARENT tags
                    run = current_para.add_run(clean_part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    
                    # Check parents for formatting
                    parent_tags = [p.name for p in element.parents]
                    parent_styles = [str(p.get('style', '')) for p in element.parents]
                    
                    # GREEN LOGIC: <ins> OR <span> with color:green
                    if 'ins' in parent_tags or any('color:green' in s or 'color: #008000' in s for s in parent_styles):
                        run.font.color.rgb = RGBColor(0, 100, 0)
                        run.font.bold = True
                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    
                    # RED LOGIC: <del> OR <span> with color:red
                    elif 'del' in parent_tags or 'strike' in parent_tags or any('color:red' in s or 'color: #800000' in s for s in parent_styles):
                        run.font.color.rgb = RGBColor(180, 0, 0)
                        run.font.strike = True

                # If we split, add a new paragraph (except for the last part)
                if i < len(parts) - 1:
                    current_para = doc.add_paragraph()
        
        # 2. If it's a Tag, dive deeper (Recursion)
        else:
            for child in element.children:
                process_element(child)

    # Start the crawl
    process_element(soup)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. MAIN UI ---
st.title("⚖️ LexRedline")
st.markdown("### Professional Document Comparison")

col1, col2 = st.columns(2)
with col1:
    f1 = st.file_uploader("Original Version", type=["pdf", "docx"], key="f1")
with col2:
    f2 = st.file_uploader("New Version", type=["pdf", "docx"], key="f2")

if st.button("Generate & Convert", type="primary"):
    if f1 and f2:
        with st.spinner("Analyzing..."):
            # 1. Get Text
            t1 = get_clean_text(f1)
            t2 = get_clean_text(f2)
            
            # 2. Compare
            # We explicitly convert Markdown -> HTML here to catch those <span> tags
            raw_html, display_html = compare_html(t1[:30000], t2[:30000])

            # 3. Generate Word Doc
            docx_file = create_docx(raw_html)

            # 4. Preview
            st.divider()
            st.subheader("Preview")
            st.markdown(f'<div class="diff-box">{display_html}</div>', unsafe_allow_html=True)
            
            # 5. Download
            st.success("✅ Analysis Complete.")
            st.download_button(
                label="📥 Download as Word (.docx)",
                data=docx_file,
                file_name="Legal_Redline.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:
        st.warning("Please upload both documents.")